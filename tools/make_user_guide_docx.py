"""Render the non-technical user guide to a Word document.

    python tools/make_user_guide_docx.py   ->   USER_GUIDE.docx

Written for someone who has never used the app and does not write code: how to
start it, how to simulate an SU email, what each of the three bundled scenarios
proves, how to read a flagged field, and how to change the rules themselves.

Check counts quoted in the text are the demo-mode (no API key) values, verified
against a real run. With a live API key the T3 numbers can shift by one, because
the model may read a borderline field more or less confidently - so T3 is
described by behaviour rather than by an exact total.
"""

from __future__ import annotations

import subprocess
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "USER_GUIDE.docx"

ACCENT = RGBColor(0x1F, 0x4E, 0x79)
MUTED = RGBColor(0x5A, 0x5A, 0x5A)
GREEN = RGBColor(0x1B, 0x6B, 0x3A)
RED = RGBColor(0xA4, 0x1E, 0x22)
CODEBG = "F4F6F8"
BOXBG = "F4F8FC"
RULE = "C9D6E4"

BODY = 10.5


# ----------------------------------------------------------------------------------
# low-level helpers
# ----------------------------------------------------------------------------------

def shade(cell, fill: str) -> None:
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(el)


def cell_border(cell, color: str) -> None:
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), color)
        borders.append(el)
    cell._tc.get_or_add_tcPr().append(borders)


def bottom_rule(paragraph, color: str, size: str = "8") -> None:
    pbdr = OxmlElement("w:pBdr")
    el = OxmlElement("w:bottom")
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), size)
    el.set(qn("w:space"), "1")
    el.set(qn("w:color"), color)
    pbdr.append(el)
    paragraph._p.get_or_add_pPr().append(pbdr)


def runs(paragraph, segments, size=BODY, color=None, mono=False):
    """segments: list of (text, style) where style contains 'b', 'i', 'c' (code)."""
    for text, style in segments:
        r = paragraph.add_run(text)
        r.bold = "b" in style
        r.italic = "i" in style
        if "c" in style or mono:
            r.font.name = "Consolas"
            r.font.size = Pt(size - 1)
        else:
            r.font.size = Pt(size)
        if color is not None:
            r.font.color.rgb = color
    return paragraph


def para(doc, segments, *, size=BODY, after=6, before=0, color=None, style=None):
    p = doc.add_paragraph(style=style)
    pf = p.paragraph_format
    pf.space_after = Pt(after)
    pf.space_before = Pt(before)
    runs(p, segments, size=size, color=color)
    return p


def h1(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = ACCENT
    bottom_rule(p, RULE, size="4")
    return p


def h2(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = ACCENT
    return p


def bullet(doc, segments, size=BODY):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    runs(p, segments, size=size)
    return p


def code_block(doc, lines: list[str], width_in: float = 6.6):
    t = doc.add_table(rows=1, cols=1)
    t.autofit = False
    c = t.rows[0].cells[0]
    c.width = Inches(width_in)
    shade(c, CODEBG)
    cell_border(c, RULE)
    c.text = ""
    for i, line in enumerate(lines):
        p = c.paragraphs[0] if i == 0 else c.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(line)
        r.font.name = "Consolas"
        r.font.size = Pt(9)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def callout(doc, segments, fill=BOXBG, width_in: float = 6.6):
    t = doc.add_table(rows=1, cols=1)
    t.autofit = False
    c = t.rows[0].cells[0]
    c.width = Inches(width_in)
    shade(c, fill)
    cell_border(c, RULE)
    c.text = ""
    p = c.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    runs(p, segments, size=BODY - 0.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def grid(doc, headers: list[str], rows: list[list[str]], widths: list[float]):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.autofit = False
    for i, head in enumerate(headers):
        c = t.rows[0].cells[i]
        c.width = Inches(widths[i])
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        shade(c, "1F4E79")
        c.text = ""
        p = c.paragraphs[0]
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.space_before = Pt(1)
        r = p.add_run(head)
        r.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            c = cells[i]
            c.width = Inches(widths[i])
            c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            if ri % 2:
                shade(c, "F7FAFC")
            c.text = ""
            p = c.paragraphs[0]
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.space_before = Pt(1)
            mono = val.startswith("`") and val.endswith("`")
            r = p.add_run(val.strip("`"))
            r.font.size = Pt(9.5)
            if mono:
                r.font.name = "Consolas"
                r.font.size = Pt(9)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


# ----------------------------------------------------------------------------------
# document
# ----------------------------------------------------------------------------------

def build() -> None:
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(BODY)
    st.paragraph_format.space_after = Pt(6)

    s = doc.sections[0]
    s.page_width, s.page_height = Inches(8.5), Inches(11)
    s.left_margin = s.right_margin = Inches(0.95)
    s.top_margin = Inches(0.8)
    s.bottom_margin = Inches(0.7)

    # ---- title -------------------------------------------------------------------
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("How to Use the Document Verification Assistant")
    r.bold = True
    r.font.size = Pt(19)
    r.font.color.rgb = ACCENT

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    runs(p, [(
        "A hands-on guide for the CG desk. No technical background needed - if you can "
        "open a folder and edit a text file, you can run and change everything here.", "i",
    )], size=10, color=MUTED)
    bottom_rule(p, "1F4E79", size="8")

    # ---- mental model ------------------------------------------------------------
    h1(doc, "First: the mental model")
    para(doc, [(
        "Think of it as a new assistant on the CG desk. An email lands, the assistant opens "
        "the attachment, checks each field against a checklist, and writes the reply - then "
        "hands it to you. You read it and hit send. ", ""),
        ("The assistant never sends anything itself.", "b"),
    ])
    para(doc, [
        ("That checklist is a plain text file. ", ""),
        ("It is the only thing that decides pass or fail", "b"),
        (" - not the AI, not the code. That matters because you can read it, and you can "
         "change it, without a developer.", ""),
    ])

    # ---- starting ----------------------------------------------------------------
    h1(doc, "Before you start: opening the app")
    para(doc, [("Open a terminal in the project folder and run these two lines. The first "
                "is only needed the very first time.", "")])
    code_block(doc, [
        "pip install -r requirements.txt",
        "python -m streamlit run app.py",
    ])
    para(doc, [("Your browser opens the app automatically. Click the first tab, ", ""),
               ("📬 Verify (Part 2)", "b"), (". That is the CG desk screen.", "")])
    callout(doc, [
        ("No API key? It still works. ", "b"),
        ("The app runs in demo mode using recorded readings of the bundled documents, and "
         "says so on screen. Everything in this guide works either way.", ""),
    ])

    # ---- 1 simulate --------------------------------------------------------------
    h1(doc, "1. Simulating an email arriving")
    para(doc, [
        ("There is no real mailbox. Instead there is a folder the app watches, called ", ""),
        ("su_inbox", "c"),
        (". Dropping a file into it is exactly like an email landing.", ""),
    ])
    para(doc, [("You do not have to do that by hand. On the ", ""),
               ("📬 Verify (Part 2)", "b"), (" tab:", "")])
    bullet(doc, [("Pick one of the three bundled emails from the dropdown", "")])
    bullet(doc, [("Click ", ""), ("✉️ Simulate this SU email arriving", "b")])
    para(doc, [(
        "That is it. The assistant wakes up, reads the attachment, checks it, and drafts a "
        "reply - all while you watch.", "")])
    para(doc, [
        ("If you would rather do it manually, drag any PDF or image straight into the ", ""),
        ("su_inbox", "c"), (" folder and click ", ""), ("📥 Check inbox now", "b"), (".", ""),
    ])

    # ---- 2 happy path ------------------------------------------------------------
    h1(doc, "2. The happy path - everything correct")
    para(doc, [("Pick ", ""), ("su_email_1_clean_invoice.json", "c"),
               (" and click Simulate.", "")])
    para(doc, [("You should see a green banner: ", ""),
               ("“All 11 checks matched.”", "b"),
               (" Every row in the table shows a green tick. Scroll down and the draft reply "
                "says ", ""), ("APPROVED", "b"), (".", "")], color=None)
    para(doc, [("Click ", ""), ("📤 Send reply (as CG)", "b"),
               (" and you are done. At that moment the document also becomes searchable in "
                "the data lake.", "")])

    # ---- 3 wrong -----------------------------------------------------------------
    h1(doc, "3. The path with details wrong")
    para(doc, [("Two flavours, deliberately different.", "")])

    h2(doc, "One thing clearly wrong")
    para(doc, [("Pick ", ""), ("su_email_2_hs_mismatch_bl.json", "c"), (". Red banner: ", ""),
               ("“1 of 8 checks need attention.”", "b"),
               (" The HS code on the document reads ", ""), ("1006.40", "b"),
               (", but this customer requires ", ""), ("1006.30", "b"),
               (". In real life that single wrong digit means a customs hold at Rotterdam.",
                "")])

    h2(doc, "Several things missing")
    para(doc, [("Pick ", ""), ("su_email_3_incomplete_invoice.json", "c"),
               (". Four fields are simply not printed on that document - the two ports, the "
                "currency and the total. The banner shows the missing and uncertain counts, "
                "and critically: ", ""),
               ("it still refuses to approve.", "b"),
               (" “Could not read it” is never treated as “it is fine.”",
                "")])

    # ---- 4 why flagged -----------------------------------------------------------
    h1(doc, "4. Checking why something was flagged")
    para(doc, [("This is the ", ""), ("🚩 Discrepancy detail", "b"),
               (" section, and it is the part worth showing people. Click a flagged field "
                "and you get two panels side by side:", "")])
    grid(
        doc,
        ["Found on the document", "What the customer requires"],
        [[
            "1006.40\nConfidence 0.95 (green) · quoted: “HS CODE: 1006.40”",
            "1006.30\n“HS code must be 1006.30 - semi-milled or wholly milled rice. Any "
            "other code risks a customs hold at Rotterdam.”",
        ]],
        [3.3, 3.3],
    )
    para(doc, [("Three things to notice, because they are what make it trustworthy:", "")])
    bullet(doc, [("The quote. ", "b"),
                 ("It shows the exact text it read off the page. You can check it in two "
                  "seconds without opening the PDF.", "")])
    bullet(doc, [("The confidence. ", "b"),
                 ("0.95 in green means it read that clearly. So this is a confident catch, "
                  "not a guess.", "")])
    bullet(doc, [("The reason in plain English. ", "b"),
                 ("The rule explains itself - that sentence comes from the rules file, "
                  "written by a human.", "")])
    para(doc, [(
        "The draft reply below lists exactly this: the field, what was sent, and what was "
        "expected. Which is what CG types out by hand today.", "")])

    # ---- 5 rules -----------------------------------------------------------------
    h1(doc, "5. Changing the rules")
    para(doc, [("The file is ", ""), ("rules/sunpeak_foods.json", "c"),
               (" - one file per customer. Each rule is five lines:", "")])
    code_block(doc, [
        "{",
        '  "field":       "hs_code",',
        '  "check":       "digits",',
        '  "expected":    "1006.30",',
        '  "label":       "HS code must be 1006.30 - any other code risks a customs hold.",',
        '  "applies_to":  ["*"]',
        "}",
    ])
    bullet(doc, [("field", "c"), (" - which box on the document: ", ""),
                 ("consignee, hs_code, incoterm, currency, origin_port, destination_port, "
                  "total_amount, gross_weight_kg, carrier, vessel_name", "c"), (" and more.",
                                                                               "")])
    bullet(doc, [("expected", "c"), (" - what it should say.", "")])
    bullet(doc, [("label", "c"), (" - your plain-English explanation. This is what CG reads "
                                  "in the app ", ""), ("and", "i"),
                 (" what goes into the email to the supplier, so write it like you are "
                  "explaining it to them.", "")])
    bullet(doc, [("applies_to", "c"), (" - ", ""), ('["*"]', "c"),
                 (" for every document, or ", ""), ('["commercial_invoice"]', "c"),
                 (" to check invoices only.", "")])
    bullet(doc, [("check", "c"), (" - how strictly to compare:", "")])
    grid(
        doc,
        ["check", "Meaning", "Use it for"],
        [
            ["`equals`", "Must match exactly", "CIF, USD"],
            ["`contains`", "Expected appears anywhere inside",
             "Rotterdam matching “Rotterdam, Netherlands”"],
            ["`name`", "Ignores punctuation and capitals",
             "Sunpeak Foods BV matching “Sunpeak Foods B.V.”"],
            ["`digits`", "Compares digits only", "1006.30 matching “100630”"],
        ],
        [1.0, 2.5, 3.1],
    )

    h2(doc, "How to update")
    para(doc, [("Edit the file and save it. ", ""), ("No restart needed", "b"),
               (" - the assistant re-reads the file every single time it checks a document.",
                "")])
    callout(doc, [
        ("One catch that will confuse you if nobody says it: ", "b"),
        ("each email is processed once, deliberately, so a real mailbox can never "
         "double-process. After changing a rule, re-simulating the same email does nothing. "
         "Open the sidebar, expand ", ""), ("Reset", "b"), (", click ", ""),
        ("Clear verifications & inbox", "b"), (", then simulate again.", ""),
    ])
    para(doc, [("Adding a new customer is just a new file alongside it, e.g. ", ""),
               ("rules/acme_foods.json", "c"),
               (". The app currently loads Sunpeak's file specifically; pointing it at a "
                "per-customer file based on the consignee is a small change, and is listed "
                "as deliberately out of scope for this version.", "")])

    # ---- test run ----------------------------------------------------------------
    h1(doc, "A good 90-second test run")
    para(doc, [("Do this in order. The last step is the one that lands with a non-technical "
                "audience, because it shows the business owns the rules, not engineering.",
                "")])
    for i, (bold, rest) in enumerate([
        ("Clean email", " - simulate it, see all green, send the approval."),
        ("HS mismatch", " - simulate it, open the discrepancy, read found vs expected."),
        ("Change the rule", " - edit \"expected\" from 1006.30 to 1006.40, save."),
        ("Clear and re-simulate", " - sidebar Reset, then run the same mismatch email again."),
        ("Watch it pass", " - the identical document now comes back clean. Nothing about the "
                          "document changed. Only the rule did."),
    ], start=1):
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        runs(p, [(bold, "b"), (rest, "")])

    # ---- troubleshooting ---------------------------------------------------------
    h1(doc, "If something looks wrong")
    grid(
        doc,
        ["What you see", "What to do"],
        [
            ["Nothing happens when I simulate an email",
             "That email was already processed. Sidebar > Reset > Clear verifications & inbox, "
             "then simulate again."],
            ["“streamlit: command not found”",
             "Use python -m streamlit run app.py instead of the bare streamlit command."],
            ["The sidebar says DEMO MODE",
             "Normal without an API key. Everything in this guide still works; readings are "
             "replayed from a recorded run and labelled on screen."],
            ["I changed a rule and nothing changed",
             "You almost certainly re-ran an already-processed email. Clear the inbox first "
             "(see the first row)."],
            ["I want a completely fresh start",
             "Delete the file data/datalake.db. It rebuilds itself the next time you start "
             "the app."],
        ],
        [2.2, 4.4],
    )

    doc.save(OUT)
    print(f"Wrote {OUT.name} ({OUT.stat().st_size:,} bytes)")


def page_count() -> None:
    ps = (
        "$ErrorActionPreference='Stop';"
        "$w = New-Object -ComObject Word.Application; $w.Visible = $false;"
        f"$d = $w.Documents.Open('{OUT}', $false, $true);"
        "$n = $d.ComputeStatistics(2); $d.Close($false); $w.Quit(); Write-Output $n"
    )
    try:
        out = subprocess.run(
            ["powershell", "-NoProfile", "-NonInteractive", "-Command", ps],
            capture_output=True, text=True, timeout=120,
        )
        lines = (out.stdout or "").strip().splitlines()
        if lines and lines[-1].isdigit():
            print(f"Word reports {lines[-1]} page(s).")
    except Exception as exc:
        print(f"Page-count check skipped ({exc}).")


if __name__ == "__main__":
    build()
    page_count()
