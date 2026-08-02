"""Render the Part 2 PRD to a one-page Word document.

    python tools/make_prd_docx.py   ->   PRD_Part2.docx

The brief asks for the PRD as a PDF or Google Doc. This .docx is the Google Doc
route (upload to Drive and it opens as one), and it is the editable format a
reviewer can comment in. PRD_Part2.md remains the source of truth for the text;
tools/make_prd_pdf.py renders the same content as PDF.

Word has no page-count API before rendering, so the one-page constraint is
verified after the fact with Word COM in `verify_page_count`. "Max 1 page" is a
scored scope-discipline constraint, not a formatting preference.
"""

from __future__ import annotations

import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "PRD_Part2.docx"

ACCENT = RGBColor(0x1F, 0x4E, 0x79)
MUTED = RGBColor(0x5A, 0x5A, 0x5A)
BOXBG = "F4F8FC"
RULE = "C9D6E4"

BODY_PT = 9
SMALL_PT = 7.5


def shade(cell, fill: str) -> None:
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(el)


def cell_border(cell, color: str) -> None:
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), color)
        borders.append(el)
    cell._tc.get_or_add_tcPr().append(borders)


def bottom_rule(paragraph, color: str, size: str = "8") -> None:
    pbdr = OxmlElement("w:pBdr")
    el = OxmlElement("w:bottom")
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), size)
    el.set(qn("w:space"), "1")
    el.set(qn("w:color"), color)
    pbdr.append(el)
    paragraph._p.get_or_add_pPr().append(pbdr)


def runs(paragraph, segments, size: float = BODY_PT, color: RGBColor | None = None):
    """segments: list of (text, style) where style is '', 'b', 'i' or 'bi'."""
    for text, style in segments:
        r = paragraph.add_run(text)
        r.font.size = Pt(size)
        r.bold = "b" in style
        r.italic = "i" in style
        if color is not None:
            r.font.color.rgb = color
    return paragraph


def para(doc, segments, *, size=BODY_PT, after=3, before=0, justify=True, color=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(after)
    pf.space_before = Pt(before)
    pf.line_spacing = 1.0
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    runs(p, segments, size=size, color=color)
    return p


def heading(doc, text: str):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(1.5)
    pf.line_spacing = 1.0
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = ACCENT
    return p


def build() -> None:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(BODY_PT)
    style.paragraph_format.space_after = Pt(3)
    style.paragraph_format.line_spacing = 1.0

    s = doc.sections[0]
    s.page_width, s.page_height = Inches(8.5), Inches(11)
    s.left_margin = s.right_margin = Inches(0.7)
    s.top_margin = Inches(0.55)
    s.bottom_margin = Inches(0.5)

    content_w = 8.5 - 1.4  # 7.1"

    # ---- title -----------------------------------------------------------------
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run("PRD: SU to CG Trade Document Verification Agent")
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = ACCENT

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    runs(
        p,
        [(
            "Part 2 · Builds on the Part 1 POC: same vision extraction, same store, "
            "same analytics layer. v1.0 · 2026-08-02", "i",
        )],
        size=8,
        color=MUTED,
    )
    bottom_rule(p, "1F4E79", size="8")

    # ---- problem ---------------------------------------------------------------
    heading(doc, "Problem")
    para(doc, [(
        "Every shipment's documents are validated by a CG team member who opens each "
        "attachment, reads every field, and checks it against what the customer requires "
        "— then types out what is wrong. Rules live in people's heads, so a new hire errs "
        "for weeks, nobody can see how many documents are pending, and there is no audit "
        "trail when a dispute surfaces. The result is 2 to 4 amendment cycles per shipment "
        "at 4 to 24 hours each, with CG bandwidth capping how fast shipments clear.", "",
    )])

    # ---- personas --------------------------------------------------------------
    heading(doc, "Personas")
    t = doc.add_table(rows=1, cols=2)
    t.autofit = False
    half = Inches(content_w / 2)
    people = [
        ("Meera — CG validator (3 yrs)",
         "Checks ~40 documents a day across a dozen customers' unwritten rule sets. Judged "
         "on one thing: errors that reach the customer (a wrong HS code means a customs hold "
         "and a penalty). Cares about catching every discrepancy without re-reading every "
         "field, and never being the reason a shipment sat for a day."),
        ("Rahul — SU documentation executive",
         "His job feels done when the doc-set email is sent. Judged on dispatch speed; every "
         "amendment email is unplanned rework. Cares about knowing exactly what to fix, in "
         "one pass — not “please recheck the invoice” ping-pong."),
    ]
    for idx, (name, text) in enumerate(people):
        cell = t.rows[0].cells[idx]
        cell.width = half
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        cell.text = ""
        p0 = cell.paragraphs[0]
        p0.paragraph_format.space_after = Pt(1)
        p0.paragraph_format.line_spacing = 1.0
        runs(p0, [(name, "b")])
        p1 = cell.add_paragraph()
        p1.paragraph_format.space_after = Pt(0)
        p1.paragraph_format.line_spacing = 1.0
        p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        runs(p1, [(text, "")])

    # ---- JTBD ------------------------------------------------------------------
    heading(doc, "Jobs to be done")
    para(doc, [
        ("1. (Meera) ", "b"),
        ("When an SU document lands in my inbox, I want a field-by-field verdict against "
         "the customer's written requirements before I open the attachment, so that I spend "
         "my time only on the fields that are actually wrong.", ""),
    ], after=1.5)
    para(doc, [
        ("2. (Rahul) ", "b"),
        ("When my documents have an issue, I want an amendment request that lists each "
         "field with what I sent and what was expected, so that I can fix everything in one "
         "cycle instead of three.", ""),
    ])

    # ---- flow ------------------------------------------------------------------
    heading(doc, "The flow (bold = human touchpoint)")
    ft = doc.add_table(rows=1, cols=1)
    ft.autofit = False
    fcell = ft.rows[0].cells[0]
    fcell.width = Inches(content_w)
    shade(fcell, BOXBG)
    cell_border(fcell, RULE)
    fcell.text = ""
    fp = fcell.paragraphs[0]
    fp.paragraph_format.space_after = Pt(0)
    fp.paragraph_format.line_spacing = 1.0
    fp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    runs(fp, [
        ("SU emails the document", "b"),
        (" → agent detects it (trigger) → agent extracts fields (Part 1 vision agent) → "
         "agent compares against the customer rule set (deterministic) → agent flags each "
         "field: match / mismatch / uncertain / missing → agent drafts the reply → ", ""),
        ("CG opens the verification result", "b"), (" → ", ""),
        ("CG inspects flagged fields", "b"),
        (" (found vs expected, with the quoted evidence) → ", ""),
        ("CG edits the draft if needed", "b"), (" → ", ""),
        ("CG SENDS", "b"), (" → ", ""),
        ("SU fixes and resends", "b"),
        (" (loop), or on a clean pass the docs go to the customer.", ""),
    ], size=8.5)

    para(doc, [(
        "The three-party structure is untouched — SU sends, CG validates, the customer "
        "receives one clean set. The agent removes the reading and the typing, not the "
        "humans. It has no send capability at all.", "",
    )], before=3)

    # ---- north star ------------------------------------------------------------
    heading(doc, "North-star metric")
    para(doc, [
        ("Median turnaround: SU email arrival to CG reply sent (minutes). ", "b"),
        ("Computed from the system's own audit trail (v_verifications.turnaround_minutes), "
         "so a CG team lead can check it on Day 14 with one query against the manual "
         "baseline of hours. Guardrail so speed never buys errors: percentage of "
         "agent-approved documents later amended (target: 0).", ""),
    ])

    # ---- failure mode ----------------------------------------------------------
    heading(doc, "Failure mode, and how it is stopped")
    para(doc, [
        ("Worst case: a false approval", "b"),
        (" — the agent shows a wrong field as matched, Meera trusts the green tick, and a "
         "bad document reaches customs. Stopped four ways: ", ""),
        ("(1)", "b"),
        (" verdicts are deterministic rule checks against a written rule set, so no model "
         "judgment can reason a mismatch away; ", ""),
        ("(2)", "b"), (" any field below the confidence bar is ", ""),
        ("uncertain", "i"),
        (", and uncertain blocks approval exactly like a mismatch — the agent never "
         "silently approves what it could not read; ", ""),
        ("(3)", "b"),
        (" every verdict carries the verbatim evidence quote, so CG can spot-check any "
         "field in seconds; ", ""),
        ("(4)", "b"),
        (" the agent cannot send — an approval only goes out through Meera's button, and "
         "the reply is rendered from the recorded check table, so it cannot claim anything "
         "the comparator did not find.", ""),
    ])

    # ---- footer ----------------------------------------------------------------
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    bottom_rule(p, RULE, size="4")

    para(doc, [(
        "Out of scope this iteration, deliberately: real email integration (the trigger is "
        "a watched folder, per the brief), multi-customer rule packs, rule learning from "
        "amendment history, and cross-document consistency (invoice vs B/L). The last two "
        "are already named as Iteration 2 in the Part 1 PRD.", "i",
    )], size=SMALL_PT, color=MUTED, after=0)

    doc.save(OUT)
    print(f"Wrote {OUT.name} ({OUT.stat().st_size:,} bytes)")


def verify_page_count() -> None:
    """Ask Word for the real rendered page count. Windows-only; skipped elsewhere."""
    ps = (
        "$ErrorActionPreference='Stop';"
        "$w = New-Object -ComObject Word.Application;"
        "$w.Visible = $false;"
        f"$d = $w.Documents.Open('{OUT}', $false, $true);"
        "$n = $d.ComputeStatistics(2);"
        "$d.Close($false); $w.Quit();"
        "Write-Output $n"
    )
    try:
        out = subprocess.run(
            ["powershell", "-NoProfile", "-NonInteractive", "-Command", ps],
            capture_output=True, text=True, timeout=120,
        )
    except Exception as exc:
        print(f"Page-count check skipped ({exc}).")
        return

    pages = (out.stdout or "").strip().splitlines()
    if not pages or not pages[-1].isdigit():
        print(f"Page-count check unavailable: {(out.stderr or '').strip()[:160]}")
        return

    n = int(pages[-1])
    print(f"Word reports {n} page(s).")
    if n != 1:
        print(f"ERROR: PRD must be exactly 1 page, got {n}.", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    build()
    verify_page_count()
